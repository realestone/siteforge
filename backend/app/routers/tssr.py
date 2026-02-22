import io
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.database import get_db
from app.models.photo import ProjectPhoto
from app.models.project import Project
from app.models.tssr import ProjectTSSR
from app.schemas.tssr import TSSRInput
from app.tssr_template_generator import generate_tssr_template
from app.tssr_template_map import TEMPLATE_MAP

router = APIRouter(prefix="/api/projects/{project_id}/tssr", tags=["tssr"])


@router.get("")
async def get_tssr(project_id: uuid.UUID, db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(ProjectTSSR).where(ProjectTSSR.project_id == project_id)
    )
    tssr = result.scalar_one_or_none()
    if not tssr:
        raise HTTPException(404, "TSSR not found for this project")
    return _tssr_to_dict(tssr)


@router.put("")
async def update_tssr(
    project_id: uuid.UUID, body: TSSRInput, db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(ProjectTSSR).where(ProjectTSSR.project_id == project_id)
    )
    tssr = result.scalar_one_or_none()
    if not tssr:
        raise HTTPException(404, "TSSR not found for this project")

    data = body.model_dump(by_alias=False)
    # Convert sector_data to plain dicts for JSONB
    data["sector_data"] = [
        s.model_dump(by_alias=False) if hasattr(s, "model_dump") else s
        for s in body.sector_data
    ]
    data["config_string"] = data.pop("config", "")
    # Convert revision_history entries to plain dicts for JSONB
    if "revision_history" in data:
        data["revision_history"] = [
            e.model_dump() if hasattr(e, "model_dump") else e
            for e in data["revision_history"]
        ]

    for field, value in data.items():
        if hasattr(tssr, field):
            setattr(tssr, field, value)

    await db.commit()
    await db.refresh(tssr)

    # TODO: trigger dependency engine recalculation here
    return _tssr_to_dict(tssr)


@router.get("/export")
async def export_tssr(
    project_id: uuid.UUID,
    format: str = Query("legacy", alias="format"),
    as_built: bool = Query(False, alias="as_built"),
    db: AsyncSession = Depends(get_db),
):
    """Export TSSR as filled .docx.

    Query params:
        format: "legacy" (fill original OneCo template) or "modern" (generate new template)
        as_built: if true, generates as-built variant with deviations and as-built photos
    """
    # Get project info
    proj_result = await db.execute(select(Project).where(Project.id == project_id))
    project = proj_result.scalar_one_or_none()
    if not project:
        raise HTTPException(404, "Project not found")

    # Get TSSR data
    result = await db.execute(
        select(ProjectTSSR).where(ProjectTSSR.project_id == project_id)
    )
    tssr = result.scalar_one_or_none()
    if not tssr:
        raise HTTPException(404, "TSSR not found for this project")

    # Get project photos for embedding — planning phase for regular, both for as-built
    photo_query = (
        select(ProjectPhoto)
        .where(ProjectPhoto.project_id == project_id)
        .where(ProjectPhoto.section != "unsorted")
        .order_by(ProjectPhoto.section, ProjectPhoto.sort_order)
    )
    if not as_built:
        photo_query = photo_query.where(ProjectPhoto.phase == "planning")
    photo_result = await db.execute(photo_query)
    all_photos = list(photo_result.scalars().all())

    planning_photos = [p for p in all_photos if (p.phase or "planning") == "planning"]
    as_built_photos = (
        [p for p in all_photos if p.phase == "as_built"] if as_built else []
    )

    # Increment as-built version if exporting as-built
    if as_built:
        project.as_built_tssr_version += 1
        await db.commit()

    if format == "modern":
        return _export_modern(
            tssr,
            project,
            planning_photos,
            as_built=as_built,
            as_built_photos=as_built_photos,
        )

    return _export_legacy(
        tssr,
        project,
        planning_photos,
        as_built=as_built,
        as_built_photos=as_built_photos,
    )


def _export_modern(
    tssr: ProjectTSSR,
    project,
    photos: list[ProjectPhoto],
    as_built: bool = False,
    as_built_photos: list[ProjectPhoto] | None = None,
) -> StreamingResponse:
    """Generate a modern TSSR template filled with project data."""
    data = {
        "site_name": tssr.site_name,
        "site_id": tssr.site_id,
        "site_model": tssr.site_model,
        "site_type": tssr.site_type,
        "customer": tssr.customer,
        "site_owner": tssr.site_owner,
        "site_category": tssr.site_category,
        "site_owner_offer": tssr.site_owner_offer,
        "montasjeunderlag": tssr.montasjeunderlag,
        "sart": tssr.sart,
        "veiviser": tssr.veiviser,
        "rfsr_rnp": tssr.rfsr_rnp,
        "guideline_version": tssr.guideline_version,
        "veiviser_comments": tssr.veiviser_comments,
        "iloq_required": tssr.iloq_required,
        "iloq_details": tssr.iloq_details,
        "access_instructions": tssr.access_instructions,
        "crane_needed": tssr.crane_needed,
        "tssr_alignment": tssr.tssr_alignment,
        "tssr_alignment_comments": tssr.tssr_alignment_comments,
    }
    buf = generate_tssr_template(data)

    # Insert planned works and photos into the generated document
    if tssr.planned_works or photos or as_built:
        doc = Document(buf)
        _insert_planned_works_into_doc(doc, tssr.planned_works)
        if photos:
            _insert_photos_into_doc(doc, photos)
        if as_built:
            _insert_as_built_section(doc, tssr, as_built_photos or [])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

    site_id = tssr.site_id or project.site_id or "export"
    if as_built:
        ver = project.as_built_tssr_version
        filename = f"{site_id}_TSSR_AsBuilt_v{ver:02d}.docx"
    else:
        filename = f"TSSR_{site_id}_modern.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_legacy(
    tssr: ProjectTSSR,
    project,
    photos: list[ProjectPhoto],
    as_built: bool = False,
    as_built_photos: list[ProjectPhoto] | None = None,
) -> StreamingResponse:
    """Fill the original OneCo .docx template with project data."""
    # Resolve template path
    template_path = Path(settings.tssr_template_path)
    if not template_path.is_absolute():
        template_path = Path(__file__).resolve().parent.parent.parent / template_path
    if not template_path.exists():
        raise HTTPException(500, f"TSSR template not found: {template_path}")

    # Load template
    doc = Document(str(template_path))

    # Build a lookup of all SDTs in the document, keyed by tag
    sdt_by_tag: dict[str, list] = {}
    for sdt in doc.element.findall(".//" + qn("w:sdt")):
        tag_el = sdt.find(".//" + qn("w:tag"))
        if tag_el is not None:
            tag_val = tag_el.get(qn("w:val"))
            if tag_val:
                sdt_by_tag.setdefault(tag_val, []).append(sdt)

    # Process each mapping entry
    for entry in TEMPLATE_MAP:
        db_field = entry["db_field"]
        value = getattr(tssr, db_field, "")

        # Check condition (e.g. iloq_details only if iloq_required)
        condition_field = entry.get("condition_field")
        if condition_field and not getattr(tssr, condition_field, False):
            value = ""

        # Add prefix if specified
        prefix = entry.get("prefix", "")
        text_value = (
            f"{prefix}{value}"
            if value
            else (prefix.rstrip(": ") + ":" if prefix else "")
        )

        entry_type = entry["type"]

        if entry_type in ("text", "combobox", "dropdown"):
            # SDT-based: find by tag and set content
            tag = entry["tag"]
            sdts = sdt_by_tag.get(tag, [])
            for sdt in sdts:
                _set_sdt_value(sdt, str(value) if value else "")

        elif entry_type == "plain_cell":
            # Positional: set cell text directly
            table_idx = entry["table"]
            row_idx = entry["row"]
            col_idx = entry["col"]
            if table_idx < len(doc.tables):
                _set_plain_cell(doc.tables[table_idx], row_idx, col_idx, text_value)

    # Insert planned works
    _insert_planned_works_into_doc(doc, tssr.planned_works)

    # Insert photos
    if photos:
        _insert_photos_into_doc(doc, photos)

    # Insert as-built content
    if as_built:
        _insert_as_built_section(doc, tssr, as_built_photos or [])

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    site_id = tssr.site_id or project.site_id or "export"
    if as_built:
        ver = project.as_built_tssr_version
        filename = f"{site_id}_TSSR_AsBuilt_v{ver:02d}.docx"
    else:
        filename = f"TSSR_{site_id}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )


# ── Photo category labels (mirror frontend photo-categories.ts) ──────

PHOTO_SECTION_LABELS = {
    "site_overview": "§1 — Site Overview",
    "delivery_access": "§1.1 — Delivery & Access",
    "hse_illustration": "§1.2 — HSE Illustration",
    "cable_route": "§3.5 — Cable Route",
    "power_diagram": "§3.6 — Power Diagram",
    "radio_plan_screenshot": "§4.3 — Radio Plan Screenshot",
    "effekt_screenshot": "§4.3 — Effekt Screenshot",
    "antenna_azimuth": "§5.1 — Antenna Azimuth",
    "antenna_placement": "§5.2 — Antenna Placement",
    "equipment_room": "§5.3 — Equipment Room",
    "site_plan": "§5.4 — Site Plan",
    "structural_calc": "§5.5 — Structural Calculation",
    "building_photos": "Appendix — Building Photos",
    "detail_photos": "Appendix — Detail Photos",
    "other": "Other Photos",
}

# Ordered list of sections for export
PHOTO_SECTION_ORDER = [
    "site_overview",
    "delivery_access",
    "hse_illustration",
    "cable_route",
    "power_diagram",
    "radio_plan_screenshot",
    "effekt_screenshot",
    "antenna_azimuth",
    "antenna_placement",
    "equipment_room",
    "site_plan",
    "structural_calc",
    "building_photos",
    "detail_photos",
    "other",
]

# Max image width in the document (A4 minus margins ≈ 17cm)
MAX_IMAGE_WIDTH_CM = 16.0


def _render_annotations(image_path: Path, annotations: list) -> io.BytesIO:
    """Render annotations onto a photo using Pillow, return JPEG buffer."""
    from PIL import Image, ImageDraw, ImageFont

    img = Image.open(image_path).convert("RGB")
    draw = ImageDraw.Draw(img)

    # Try to get a font for text annotations
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 24)
    except Exception:
        font = ImageFont.load_default()

    color_map = {
        "red": (255, 0, 0),
        "yellow": (255, 255, 0),
        "blue": (0, 100, 255),
        "white": (255, 255, 255),
        "black": (0, 0, 0),
    }

    for ann in annotations:
        ann_type = ann.get("type", "")
        color = color_map.get(ann.get("color", "red"), (255, 0, 0))
        points = ann.get("points", [])
        width = 3

        if ann_type == "arrow" and len(points) >= 2:
            x1, y1 = points[0]["x"], points[0]["y"]
            x2, y2 = points[1]["x"], points[1]["y"]
            draw.line([(x1, y1), (x2, y2)], fill=color, width=width)
            # Simple arrowhead
            import math

            angle = math.atan2(y2 - y1, x2 - x1)
            head_len = 15
            for offset in [2.5, -2.5]:
                hx = x2 - head_len * math.cos(angle + offset)
                hy = y2 - head_len * math.sin(angle + offset)
                draw.line([(x2, y2), (hx, hy)], fill=color, width=width)

        elif ann_type == "line" and len(points) >= 2:
            x1, y1 = points[0]["x"], points[0]["y"]
            x2, y2 = points[1]["x"], points[1]["y"]
            draw.line([(x1, y1), (x2, y2)], fill=color, width=width)

        elif ann_type == "rectangle" and len(points) >= 2:
            x1, y1 = points[0]["x"], points[0]["y"]
            x2, y2 = points[1]["x"], points[1]["y"]
            draw.rectangle([(x1, y1), (x2, y2)], outline=color, width=width)

        elif ann_type == "circle" and len(points) >= 2:
            x1, y1 = points[0]["x"], points[0]["y"]
            x2, y2 = points[1]["x"], points[1]["y"]
            draw.ellipse([(x1, y1), (x2, y2)], outline=color, width=width)

        elif ann_type == "text" and len(points) >= 1:
            x, y = points[0]["x"], points[0]["y"]
            label = ann.get("label", "")
            if label:
                draw.text((x, y), label, fill=color, font=font)

        elif ann_type == "measure" and len(points) >= 2:
            x1, y1 = points[0]["x"], points[0]["y"]
            x2, y2 = points[1]["x"], points[1]["y"]
            draw.line([(x1, y1), (x2, y2)], fill=color, width=width)
            dist = ann.get("measureDistance", "")
            if dist:
                mx, my = (x1 + x2) / 2, (y1 + y2) / 2
                draw.text((mx, my - 12), f"{dist}m", fill=color, font=font)

    buf = io.BytesIO()
    img.save(buf, "JPEG", quality=90)
    buf.seek(0)
    return buf


def _insert_planned_works_into_doc(doc: Document, planned_works: dict | None) -> None:
    """Insert planned works sections into the document after site data tables."""
    if not planned_works:
        return
    sections = planned_works.get("sections", [])
    if not sections:
        return

    doc.add_page_break()
    heading = doc.add_heading("Description of Planned Works", level=1)
    heading.runs[0].font.size = Pt(16)

    for sec in sections:
        title = sec.get("title", "")
        items = sec.get("items", [])
        if not title or not items:
            continue

        doc.add_heading(title, level=2)

        for item in items:
            text = item.get("text", "").strip()
            # For manual-field-only items, compose text from fields
            if not text:
                fields = item.get("manualFields", [])
                parts = []
                for f in fields:
                    val = f.get("value", "").strip()
                    if val:
                        label = f.get("label", "")
                        parts.append(f"{label}: {val}" if label else val)
                text = "; ".join(parts)
            if not text:
                continue

            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(text)
            run.font.size = Pt(10)

            # Add derivation as small italic note
            derivation = item.get("derivation", "")
            if derivation:
                note_run = para.add_run(f"  [{derivation}]")
                note_run.font.size = Pt(8)
                note_run.font.italic = True
                note_run.font.color.rgb = None  # gray doesn't work without RGBColor


def _insert_photos_into_doc(doc: Document, photos: list[ProjectPhoto]) -> None:
    """Insert photos grouped by TSSR section at the end of the document."""
    uploads_root = Path(settings.uploads_dir)

    # Group photos by section in defined order
    by_section: dict[str, list[ProjectPhoto]] = {}
    for photo in photos:
        section = photo.section or "other"
        by_section.setdefault(section, []).append(photo)

    # Add page break before photo appendix
    doc.add_page_break()

    # Section heading
    heading = doc.add_heading("Photo Documentation", level=1)
    heading.runs[0].font.size = Pt(16)

    for section_key in PHOTO_SECTION_ORDER:
        section_photos = by_section.get(section_key)
        if not section_photos:
            continue

        label = PHOTO_SECTION_LABELS.get(section_key, section_key)
        doc.add_heading(label, level=2)

        for photo in section_photos:
            file_path = uploads_root / photo.file_path
            if not file_path.exists():
                continue

            # Render annotations onto image if present
            annotations = photo.annotations or []
            if annotations and isinstance(annotations, list) and len(annotations) > 0:
                try:
                    img_buf = _render_annotations(file_path, annotations)
                    doc.add_picture(img_buf, width=Cm(MAX_IMAGE_WIDTH_CM))
                except Exception as e:
                    print(
                        f"Failed to render annotations for {photo.original_filename}: {e}"
                    )
                    doc.add_picture(str(file_path), width=Cm(MAX_IMAGE_WIDTH_CM))
            else:
                try:
                    doc.add_picture(str(file_path), width=Cm(MAX_IMAGE_WIDTH_CM))
                except Exception as e:
                    print(f"Failed to insert photo {photo.original_filename}: {e}")
                    continue

            # Caption
            caption_text = photo.auto_filename or photo.original_filename
            if photo.caption:
                caption_text += f" — {photo.caption}"
            caption_para = doc.add_paragraph()
            run = caption_para.add_run(caption_text)
            run.font.size = Pt(9)
            run.font.italic = True
            caption_para.paragraph_format.space_after = Pt(12)


def _insert_as_built_section(
    doc: Document, tssr: ProjectTSSR, as_built_photos: list[ProjectPhoto]
) -> None:
    """Insert as-built documentation: title label, deviations, and as-built photos."""
    from datetime import date

    doc.add_page_break()

    # As-Built heading
    heading = doc.add_heading("AS-BUILT DOCUMENTATION", level=1)
    heading.runs[0].font.size = Pt(16)

    # Date
    para = doc.add_paragraph()
    run = para.add_run(f"As-Built Date: {date.today().strftime('%d.%m.%Y')}")
    run.font.size = Pt(10)

    # Deviations free text
    deviations_text = tssr.deviations_free_text
    if deviations_text:
        doc.add_heading("Deviation Notes", level=2)
        para = doc.add_paragraph()
        run = para.add_run(deviations_text)
        run.font.size = Pt(10)

    # As-built photos
    if as_built_photos:
        doc.add_heading("As-Built Photo Documentation", level=2)
        _insert_photos_into_doc(doc, as_built_photos)


def _set_sdt_value(sdt_element, value: str) -> None:
    """Set the display text of a Structured Document Tag (SDT).

    Works for TEXT, COMBOBOX, and DROPDOWN SDTs. Handles two layouts:
    1. sdtContent > w:r (run directly in content — common in table cells)
    2. sdtContent > w:p > w:r (run inside paragraph)
    Preserves the first run's formatting (rPr) and replaces its text.
    """
    content = sdt_element.find(qn("w:sdtContent"))
    if content is None:
        return

    # Case 1: Runs directly under sdtContent (no paragraph wrapper)
    direct_runs = content.findall(qn("w:r"))
    if direct_runs:
        first_run = direct_runs[0]
        for run in direct_runs[1:]:
            content.remove(run)
        t_el = first_run.find(qn("w:t"))
        if t_el is None:
            t_el = first_run.makeelement(qn("w:t"), {})
            first_run.append(t_el)
        t_el.text = value
        t_el.set(qn("xml:space"), "preserve")
        return

    # Case 2: Runs inside a paragraph
    para = content.find(qn("w:p"))
    if para is not None:
        runs = para.findall(qn("w:r"))
        if runs:
            first_run = runs[0]
            for run in runs[1:]:
                para.remove(run)
            t_el = first_run.find(qn("w:t"))
            if t_el is None:
                t_el = first_run.makeelement(qn("w:t"), {})
                first_run.append(t_el)
            t_el.text = value
            t_el.set(qn("xml:space"), "preserve")
        else:
            run = para.makeelement(qn("w:r"), {})
            t_el = run.makeelement(qn("w:t"), {})
            t_el.text = value
            t_el.set(qn("xml:space"), "preserve")
            run.append(t_el)
            para.append(run)


def _set_plain_cell(table, row: int, col: int, value: str) -> None:
    """Set a plain table cell's text, preserving the first run's formatting."""
    if row >= len(table.rows) or col >= len(table.rows[row].cells):
        return
    cell = table.rows[row].cells[col]
    if cell.paragraphs and cell.paragraphs[0].runs:
        first_run = cell.paragraphs[0].runs[0]
        font_props = {
            "bold": first_run.font.bold,
            "italic": first_run.font.italic,
            "size": first_run.font.size,
            "name": first_run.font.name,
            "color_rgb": first_run.font.color.rgb
            if first_run.font.color and first_run.font.color.rgb
            else None,
        }
        for p in cell.paragraphs:
            p.clear()
        run = cell.paragraphs[0].add_run(str(value))
        run.font.bold = font_props["bold"]
        run.font.italic = font_props["italic"]
        run.font.size = font_props["size"]
        run.font.name = font_props["name"]
        if font_props["color_rgb"]:
            run.font.color.rgb = font_props["color_rgb"]
    else:
        cell.text = str(value)


def _tssr_to_dict(tssr: ProjectTSSR) -> dict:
    return {
        "id": str(tssr.id),
        "projectId": str(tssr.project_id),
        # Site Identity
        "siteId": tssr.site_id,
        "siteName": tssr.site_name,
        "operator": tssr.operator,
        "siteModel": tssr.site_model,
        "siteType": tssr.site_type,
        "customer": tssr.customer,
        "siteOwner": tssr.site_owner,
        # Supporting Documents
        "siteOwnerOffer": tssr.site_owner_offer,
        "montasjeunderlag": tssr.montasjeunderlag,
        "sart": tssr.sart,
        "veiviser": tssr.veiviser,
        "rfsrRnp": tssr.rfsr_rnp,
        "guidelineVersion": tssr.guideline_version,
        # Access Info
        "veiviserComments": tssr.veiviser_comments,
        "iloqRequired": tssr.iloq_required,
        "iloqDetails": tssr.iloq_details,
        # TSSR Alignment
        "tssrAlignment": tssr.tssr_alignment,
        "tssrAlignmentComments": tssr.tssr_alignment_comments,
        # Radio Configuration
        "sectors": tssr.sectors,
        "size": tssr.size,
        "config": tssr.config_string,
        "sectorData": tssr.sector_data or [],
        # Access & Logistics
        "siteCategory": tssr.site_category,
        "landlordName": tssr.landlord_name,
        "accessInstructions": tssr.access_instructions,
        "craneNeeded": tssr.crane_needed,
        # Cabinet & Power
        "cabinetType": tssr.cabinet_type,
        "acdb": tssr.acdb,
        "rectifier": tssr.rectifier,
        "earthing": tssr.earthing,
        # HSE
        "hseHazards": tssr.hse_hazards or [],
        # Building Info
        "roofType": tssr.roof_type,
        "roofMaterial": tssr.roof_material,
        "roofLoad": tssr.roof_load,
        "towerHeight": tssr.tower_height,
        # Cable Routing
        "cableLadderLength": tssr.cable_ladder_length,
        "verticalCableRoute": tssr.vertical_cable_route,
        # Antenna Mounting
        "mountType": tssr.mount_type,
        # Services
        "paintingRequired": tssr.painting_required,
        "paintingColor": tssr.painting_color,
        # Revision History
        "revisionHistory": tssr.revision_history or [],
        # Planned Works
        "plannedWorks": tssr.planned_works,
        # Other
        "additionalNotes": tssr.additional_notes,
    }
