"""
TSSR Template Generator — Modern "Site Identity & Access" page.

Generates a professional .docx template with:
  - Branded header with blue accent
  - Two-column table layout for all first-page fields
  - Structured Document Tags (SDTs) for dropdowns
  - Content controls for text inputs
  - Proper spacing, shading, and typography
  - Footer with version and date

Can be used standalone (generates blank template) or filled with project data.
"""

import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor
from docx.table import _Cell

# ── Color palette ────────────────────────────────────────────────────

BLUE_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)  # Deep blue
BLUE_LIGHT = RGBColor(0xDB, 0xEA, 0xFE)  # Light blue bg
BLUE_ACCENT = RGBColor(0x3B, 0x82, 0xF6)  # Accent blue
GRAY_900 = RGBColor(0x11, 0x18, 0x27)  # Near-black text
GRAY_700 = RGBColor(0x37, 0x41, 0x51)  # Dark gray labels
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)  # Medium gray
GRAY_400 = RGBColor(0x9C, 0xA3, 0xAF)  # Placeholder text
GRAY_200 = RGBColor(0xE5, 0xE7, 0xEB)  # Borders
GRAY_100 = RGBColor(0xF3, 0xF4, 0xF6)  # Light bg
GRAY_50 = RGBColor(0xF9, 0xFA, 0xFB)  # Very light bg
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x05, 0x96, 0x69)  # Success green
AMBER = RGBColor(0xD9, 0x77, 0x06)  # Warning amber

FONT_FAMILY = "Segoe UI"
FONT_FAMILY_MONO = "Cascadia Code"


# ── Helper functions ─────────────────────────────────────────────────


def set_cell_shading(cell: _Cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_border(cell: _Cell, **kwargs) -> None:
    """Set borders on a table cell. kwargs: top, bottom, start, end with (size, color)."""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, (size, color) in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)


def set_cell_margins(cell: _Cell, top=0, bottom=0, start=0, end=0) -> None:
    """Set inner margins (padding) on a cell in twips."""
    tc_pr = cell._element.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)


def set_cell_width(cell: _Cell, width_cm: float) -> None:
    """Set explicit cell width."""
    tc_pr = cell._element.get_or_add_tcPr()
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(int(width_cm * 567)))  # cm to twips
    width.set(qn("w:type"), "dxa")
    tc_pr.append(width)


def add_run(
    paragraph,
    text: str,
    font_name=FONT_FAMILY,
    size=Pt(11),
    color=GRAY_900,
    bold=False,
    italic=False,
) -> None:
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def remove_table_borders(table) -> None:
    """Remove all borders from a table."""
    tbl = table._element
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def add_dropdown_sdt(
    cell: _Cell,
    tag: str,
    alias: str,
    options: list[str],
    default: str = "",
    placeholder: str = "Select...",
) -> None:
    """Add a ComboBox Structured Document Tag to a cell."""
    tc = cell._element
    # Clear existing paragraphs
    for p in tc.findall(qn("w:p")):
        tc.remove(p)

    sdt = OxmlElement("w:sdt")

    # SDT Properties
    sdt_pr = OxmlElement("w:sdtPr")
    # Tag
    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdt_pr.append(tag_el)
    # Alias
    alias_el = OxmlElement("w:alias")
    alias_el.set(qn("w:val"), alias)
    sdt_pr.append(alias_el)
    # Placeholder
    ph = OxmlElement("w:placeholder")
    doc_part = OxmlElement("w:docPart")
    doc_part.set(qn("w:val"), "DefaultPlaceholder_-1854013440")
    ph.append(doc_part)
    sdt_pr.append(ph)
    # ComboBox with options
    combo = OxmlElement("w:comboBox")
    for opt in options:
        li = OxmlElement("w:listItem")
        li.set(qn("w:displayText"), opt)
        li.set(qn("w:value"), opt)
        combo.append(li)
    sdt_pr.append(combo)
    sdt.append(sdt_pr)

    # SDT Content
    sdt_content = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    # Paragraph properties for alignment
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "40")
    spacing.set(qn("w:after"), "40")
    p_pr.append(spacing)
    p.append(p_pr)
    # Run with default/placeholder text
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_font = OxmlElement("w:rFonts")
    r_font.set(qn("w:ascii"), FONT_FAMILY)
    r_font.set(qn("w:hAnsi"), FONT_FAMILY)
    r_pr.append(r_font)
    r_sz = OxmlElement("w:sz")
    r_sz.set(qn("w:val"), "22")  # 11pt
    r_pr.append(r_sz)
    if not default:
        r_color = OxmlElement("w:color")
        r_color.set(qn("w:val"), "9CA3AF")
        r_pr.append(r_color)
        r_italic = OxmlElement("w:i")
        r_pr.append(r_italic)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = default if default else placeholder
    r.append(t)
    p.append(r)
    sdt_content.append(p)
    sdt.append(sdt_content)

    tc.append(sdt)


def add_text_sdt(
    cell: _Cell,
    tag: str,
    alias: str,
    default: str = "",
    placeholder: str = "Click to enter text...",
) -> None:
    """Add a Text Structured Document Tag to a cell."""
    tc = cell._element
    for p in tc.findall(qn("w:p")):
        tc.remove(p)

    sdt = OxmlElement("w:sdt")

    # SDT Properties
    sdt_pr = OxmlElement("w:sdtPr")
    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdt_pr.append(tag_el)
    alias_el = OxmlElement("w:alias")
    alias_el.set(qn("w:val"), alias)
    sdt_pr.append(alias_el)
    # Show as bounding box
    show_as = OxmlElement("w:showingPlcHdr")
    if not default:
        sdt_pr.append(show_as)
    # Text type
    text_el = OxmlElement("w:text")
    sdt_pr.append(text_el)
    sdt.append(sdt_pr)

    # SDT Content
    sdt_content = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "40")
    spacing.set(qn("w:after"), "40")
    p_pr.append(spacing)
    p.append(p_pr)

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_font = OxmlElement("w:rFonts")
    r_font.set(qn("w:ascii"), FONT_FAMILY)
    r_font.set(qn("w:hAnsi"), FONT_FAMILY)
    r_pr.append(r_font)
    r_sz = OxmlElement("w:sz")
    r_sz.set(qn("w:val"), "22")
    r_pr.append(r_sz)
    if not default:
        r_color = OxmlElement("w:color")
        r_color.set(qn("w:val"), "9CA3AF")
        r_pr.append(r_color)
        r_italic = OxmlElement("w:i")
        r_pr.append(r_italic)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = default if default else placeholder
    r.append(t)
    p.append(r)
    sdt_content.append(p)
    sdt.append(sdt_content)

    tc.append(sdt)


# ── Field row builders ───────────────────────────────────────────────


def _label_cell(cell: _Cell, text: str, required: bool = False) -> None:
    """Format a label cell (left column)."""
    set_cell_shading(cell, "F9FAFB")
    set_cell_margins(cell, top=80, bottom=80, start=140, end=80)
    set_cell_border(cell, end=("4", "E5E7EB"))

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=Pt(10.5), color=GRAY_700, bold=True)
    if required:
        add_run(p, " *", size=Pt(10.5), color=RGBColor(0xDC, 0x26, 0x26), bold=True)


def _value_cell(cell: _Cell) -> None:
    """Format a value cell (right column) base styling."""
    set_cell_margins(cell, top=80, bottom=80, start=140, end=140)


def add_text_field(
    table,
    label: str,
    tag: str,
    alias: str,
    placeholder: str = "Enter value...",
    required: bool = False,
    default: str = "",
) -> None:
    """Add a label + text SDT row."""
    row = table.add_row()
    _label_cell(row.cells[0], label, required)
    _value_cell(row.cells[1])
    add_text_sdt(row.cells[1], tag, alias, default=default, placeholder=placeholder)


def add_dropdown_field(
    table,
    label: str,
    tag: str,
    alias: str,
    options: list[str],
    placeholder: str = "Select...",
    required: bool = False,
    default: str = "",
) -> None:
    """Add a label + dropdown SDT row."""
    row = table.add_row()
    _label_cell(row.cells[0], label, required)
    _value_cell(row.cells[1])
    add_dropdown_sdt(
        row.cells[1], tag, alias, options, default=default, placeholder=placeholder
    )


def add_multiline_field(
    table,
    label: str,
    tag: str,
    alias: str,
    placeholder: str = "Enter details...",
    default: str = "",
    rows: int = 3,
) -> None:
    """Add a label + multiline text row with extra height."""
    row = table.add_row()
    _label_cell(row.cells[0], label)
    _value_cell(row.cells[1])

    # Set minimum row height for multiline appearance
    tr = row._element
    tr_pr = tr.get_or_add_trPr()
    row_height = OxmlElement("w:trHeight")
    row_height.set(qn("w:val"), str(rows * 320))
    row_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(row_height)

    add_text_sdt(row.cells[1], tag, alias, default=default, placeholder=placeholder)


def add_section_divider(table, title: str) -> None:
    """Add a colored section divider row spanning both columns."""
    row = table.add_row()
    # Merge the two cells
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_shading(cell, "1E40AF")
    set_cell_margins(cell, top=60, bottom=60, start=140, end=140)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, title, size=Pt(9.5), color=WHITE, bold=True)


def add_toggle_field(
    table,
    label: str,
    tag: str,
    alias: str,
    detail_tag: str = "",
    detail_alias: str = "",
    detail_placeholder: str = "Enter details...",
    default_toggle: str = "",
    default_detail: str = "",
) -> None:
    """Add a Yes/No toggle field with optional detail text below."""
    # Toggle row
    add_dropdown_field(
        table,
        label,
        tag,
        alias,
        options=["Yes", "No"],
        placeholder="Select...",
        default=default_toggle,
    )

    # Detail row (conditional in practice, always shown in template)
    if detail_tag:
        row = table.add_row()
        _label_cell(row.cells[0], f"  {label} Details")
        _value_cell(row.cells[1])

        tr = row._element
        tr_pr = tr.get_or_add_trPr()
        row_height = OxmlElement("w:trHeight")
        row_height.set(qn("w:val"), "640")
        row_height.set(qn("w:hRule"), "atLeast")
        tr_pr.append(row_height)

        add_text_sdt(
            row.cells[1],
            detail_tag,
            detail_alias,
            default=default_detail,
            placeholder=detail_placeholder,
        )


def add_row_separator(table) -> None:
    """Add an empty thin separator row."""
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_shading(cell, "E5E7EB")
    tr = row._element
    tr_pr = tr.get_or_add_trPr()
    row_height = OxmlElement("w:trHeight")
    row_height.set(qn("w:val"), "30")
    row_height.set(qn("w:hRule"), "exact")
    tr_pr.append(row_height)


# ── Version options (shared) ─────────────────────────────────────────

VERSION_OPTIONS = [
    "No",
    "NA",
    "v01",
    "v02",
    "v03",
    "v04",
    "v05",
    "v06",
    "v07",
    "v08",
    "v09",
    "v10",
]

SITE_OWNER_OPTIONS = [
    "Telia Infra",
    "Telenor Infra",
    "Norkring",
    "Haugaland Kraft",
    "Lyse Fiber",
    "Broadnet",
    "Private",
    "Other",
]

BUILDING_TYPE_OPTIONS = ["Silo", "Barn", "Private House", "Factory", "Other Rooftop"]

SITE_CATEGORY_OPTIONS = ["Rooftop", "Greenfield", "Barn", "Indoor", "Tower"]

CONTRACTUAL_MODEL_OPTIONS = ["Not Applicable", "Private Site", "Coloc", "Greenfield"]


# ── Main generator ───────────────────────────────────────────────────


def generate_tssr_template(data: dict | None = None) -> io.BytesIO:
    """Generate a modern TSSR 'Site Identity & Access' .docx template.

    Args:
        data: Optional dict of field values to pre-fill. Keys match SDT tags.
              If None, generates a blank template with placeholders.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    d = data or {}
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Mm(210)  # A4
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_FAMILY
    font.size = Pt(11)
    font.color.rgb = GRAY_900

    # ── Header bar ───────────────────────────────────────────────
    # Blue accent line at top
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)

    # Set column widths
    h_left = header_table.rows[0].cells[0]
    h_right = header_table.rows[0].cells[1]
    set_cell_width(h_left, 12.0)
    set_cell_width(h_right, 5.4)

    # Blue bar under header
    for cell in header_table.rows[0].cells:
        set_cell_shading(cell, "1E40AF")
        set_cell_margins(cell, top=100, bottom=100, start=160, end=160)

    # Left: Title
    p_title = h_left.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    add_run(
        p_title, "TECHNICAL SITE SURVEY REPORT", size=Pt(16), color=WHITE, bold=True
    )

    # Right: Version + date
    p_ver = h_right.paragraphs[0]
    p_ver.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ver.paragraph_format.space_before = Pt(0)
    p_ver.paragraph_format.space_after = Pt(0)
    add_run(p_ver, "SiteForge v1.0", size=Pt(9), color=RGBColor(0xBF, 0xDB, 0xFE))
    p_ver.add_run("\n")
    add_run(
        p_ver,
        date.today().strftime("%d %b %Y"),
        size=Pt(9),
        color=RGBColor(0xBF, 0xDB, 0xFE),
    )

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)

    # ── Section title ────────────────────────────────────────────
    p_section = doc.add_paragraph()
    p_section.paragraph_format.space_before = Pt(2)
    p_section.paragraph_format.space_after = Pt(8)
    add_run(
        p_section, "Site Identity & Access", size=Pt(14), color=BLUE_PRIMARY, bold=True
    )

    # ── Main form table ──────────────────────────────────────────
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    # Set column widths via tblGrid
    tbl = table._element
    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for width in [3600, 6200]:  # twips: ~6.4cm label, ~10.9cm value
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    # ── Group 1: Site Identity ───────────────────────────────────
    add_section_divider(table, "SITE IDENTITY")

    add_text_field(
        table,
        "Site Name",
        "SiteName",
        "Site Name",
        placeholder="e.g., Oslo Sentrum Tower",
        required=True,
        default=d.get("site_name", ""),
    )

    add_text_field(
        table,
        "Site ID",
        "SiteID",
        "Site ID",
        placeholder="e.g., OSL-1234",
        required=True,
        default=d.get("site_id", ""),
    )

    add_dropdown_field(
        table,
        "Building Type",
        "BuildingType",
        "Building Type",
        options=BUILDING_TYPE_OPTIONS,
        placeholder="Select building type...",
        default=d.get("site_model", ""),
    )

    add_dropdown_field(
        table,
        "Contractual Model",
        "ContractualModel",
        "Contractual Model",
        options=CONTRACTUAL_MODEL_OPTIONS,
        placeholder="Select model...",
        default=d.get("site_type", ""),
    )

    add_dropdown_field(
        table,
        "Customer",
        "Customer",
        "Customer",
        options=["ICE", "Telenor", "Telia"],
        placeholder="Select customer...",
        required=True,
        default=d.get("customer", ""),
    )

    add_dropdown_field(
        table,
        "Site Owner",
        "SiteOwner",
        "Site Owner",
        options=SITE_OWNER_OPTIONS,
        placeholder="Select owner...",
        default=d.get("site_owner", ""),
    )

    add_dropdown_field(
        table,
        "Site Category",
        "SiteCategory",
        "Site Category",
        options=SITE_CATEGORY_OPTIONS,
        placeholder="Select category...",
        default=d.get("site_category", ""),
    )

    add_row_separator(table)

    # ── Group 2: Supporting Documents ────────────────────────────
    add_section_divider(table, "SUPPORTING DOCUMENTS")

    add_dropdown_field(
        table,
        "Site Owner Offer",
        "SiteOwnerOfferVersion",
        "Site Owner Offer Version",
        options=VERSION_OPTIONS,
        placeholder="Select version...",
        default=d.get("site_owner_offer", ""),
    )

    add_dropdown_field(
        table,
        "Montasjeunderlag",
        "MontasjeunderlagVersion",
        "Montasjeunderlag Version",
        options=VERSION_OPTIONS,
        placeholder="Select version...",
        default=d.get("montasjeunderlag", ""),
    )

    add_dropdown_field(
        table,
        "SART",
        "SARTVersion",
        "SART Version",
        options=VERSION_OPTIONS,
        placeholder="Select version...",
        default=d.get("sart", ""),
    )

    add_dropdown_field(
        table,
        "Veiviser",
        "VeiviserAvailable",
        "Veiviser Available",
        options=["Yes", "No"],
        placeholder="Select...",
        default=d.get("veiviser", ""),
    )

    add_dropdown_field(
        table,
        "RFSR / RNP",
        "RFSRVersion",
        "RFSR Version",
        options=VERSION_OPTIONS,
        placeholder="Select version...",
        default=d.get("rfsr_rnp", ""),
    )

    add_text_field(
        table,
        "Other Documents",
        "OtherSupportingDocuments",
        "Other Supporting Documents",
        placeholder="Guideline version, additional references...",
        default=d.get("guideline_version", ""),
    )

    add_row_separator(table)

    # ── Group 3: Access & Logistics ──────────────────────────────
    add_section_divider(table, "ACCESS & LOGISTICS")

    add_multiline_field(
        table,
        "Veiviser Comments",
        "VeiviserComments",
        "Veiviser Comments",
        placeholder="Access directions, key codes, contact person, parking...",
        default=d.get("veiviser_comments", ""),
        rows=3,
    )

    add_toggle_field(
        table,
        "iLOQ Required",
        "iLOQRequired",
        "iLOQ Required",
        detail_tag="iLOQDetails",
        detail_alias="iLOQ Details",
        detail_placeholder="Location, lock ID, access level...",
        default_toggle="Yes" if d.get("iloq_required") else "",
        default_detail=d.get("iloq_details", ""),
    )

    add_multiline_field(
        table,
        "Access Instructions",
        "AccessInstructions",
        "Access Instructions",
        placeholder="Detailed access instructions for site visit...",
        default=d.get("access_instructions", ""),
        rows=3,
    )

    add_toggle_field(
        table,
        "Crane Needed",
        "CraneNeeded",
        "Crane Needed",
        default_toggle="Yes" if d.get("crane_needed") else "",
    )

    add_row_separator(table)

    # ── Group 4: TSSR Alignment ──────────────────────────────────
    add_section_divider(table, "TSSR ALIGNMENT")

    add_dropdown_field(
        table,
        "TSSR Aligned",
        "TSSRAligned",
        "TSSR Aligned",
        options=["Yes", "No"],
        placeholder="Select...",
        default=d.get("tssr_alignment", ""),
    )

    add_multiline_field(
        table,
        "Alignment Comments",
        "TSSRAlignmentComments",
        "TSSR Alignment Comments",
        placeholder="Deviations, scope changes, alignment notes...",
        default=d.get("tssr_alignment_comments", ""),
        rows=3,
    )

    # ── Bottom border ────────────────────────────────────────────
    # Add a final thin blue line
    end_p = doc.add_paragraph()
    end_p.paragraph_format.space_before = Pt(6)
    end_p.paragraph_format.space_after = Pt(0)
    # Add bottom border to paragraph
    p_pr = end_p._element.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E40AF")
    p_borders.append(bottom)
    p_pr.append(p_borders)

    # ── Footer ───────────────────────────────────────────────────
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(4)
    add_run(
        p_footer,
        f"TSSR generated by SiteForge  |  {date.today().strftime('%Y-%m-%d')}  |  v1.0",
        size=Pt(8),
        color=GRAY_500,
    )

    # ── Save to buffer ───────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_tssr_template_file(
    output_path: str | Path, data: dict | None = None
) -> Path:
    """Generate template and save to disk."""
    path = Path(output_path)
    buf = generate_tssr_template(data)
    path.write_bytes(buf.getvalue())
    return path


# ── CLI entry point ──────────────────────────────────────────────────

if __name__ == "__main__":
    out = generate_tssr_template_file("TSSR_Template_Modern.docx")
    print(f"Generated: {out}")

    # Also generate a filled example
    example_data = {
        "site_name": "Kringsjaa Tower",
        "site_id": "OSL-1234",
        "site_model": "Other Rooftop",
        "site_type": "Coloc",
        "customer": "ICE",
        "site_owner": "Telia Infra",
        "site_category": "Rooftop",
        "site_owner_offer": "v02",
        "montasjeunderlag": "v01",
        "sart": "v03",
        "veiviser": "Yes",
        "rfsr_rnp": "NA",
        "guideline_version": "TSSR v6 guideline reference",
        "veiviser_comments": "Access via back door, key in lockbox #42.\nContact: Ola Nordmann +47 123 45 678",
        "iloq_required": True,
        "iloq_details": "iLOQ on main entrance and equipment room. Cylinder ID: EQ-4521",
        "access_instructions": "Drive to parking P2, walk 200m north.\nTake elevator to roof level.\nKey for rooftop door in key cabinet next to elevator.",
        "crane_needed": False,
        "tssr_alignment": "Yes",
        "tssr_alignment_comments": "Aligned with SART v03 and Montasjeunderlag v01.\nNo deviations from planned scope.",
    }
    out2 = generate_tssr_template_file("TSSR_Template_Filled.docx", example_data)
    print(f"Generated filled: {out2}")
